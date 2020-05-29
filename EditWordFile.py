#!/usr/bin/python
# coding: utf-8


from docx import Document
from convDoc2Docx import convDocToDocx as cnv

#! there are mainly three parameters in this class:
# @params in_dict is a dictionary，key: origin string，value: replaced string.
# @params origin_word_file: the path of origin docx file (or template word file).
# @params new_word_file: the path of docx file (saved file path).


class WordPro():
    def __init__(self, in_dict={"target str": "new str"}, origin_word_file='./old.docx', new_word_file='./new.docx'):
        self.replace_DICT = in_dict
        if origin_word_file.split(".")[-1] == 'docx':
            document = Document(origin_word_file)
            document = self.process(document)
            document.save(new_word_file)
        elif origin_word_file.split(".")[-1] == 'doc':
            new_docx_file = cnv(origin_word_file)
            document = Document(new_docx_file)
            document = self.process(document)
            document.save(new_word_file)
        else:
            print("input file is illegal.")

    def process(self, document):
        
        for para in document.paragraphs:
            print(para.text)

        # tables
        for table in document.tables:
            for row in range(len(table.rows)):
                for col in range(len(table.columns)):
                    for key, value in self.replace_DICT.items():
                        if key in table.cell(row, col).text:
                            print(key+"->"+value)
                            table.cell(row, col).text = table.cell(row, col).text.replace(key, value)
        # paragraphs
        for para in document.paragraphs:
            for i in range(len(para.runs)):
                for key, value in self.replace_DICT.items():
                    if key in para.runs[i].text:
                        print(key+"->"+value)
                        para.runs[i].text = para.runs[i].text.replace(key, value)
        
        for para in document.paragraphs:
            print(para.text)
        
        return document

if __name__ == '__main__':
    dict1 = {'http': 'Devin'}
    WordPro(in_dict=dict1)
    #! params origin_word_file and new_word_file is used default.
